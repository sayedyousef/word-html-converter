# integrated_converter.py
"""Integrated converter that fixes equation issues and adds anchor support."""

import mammoth
import logging
import re
import docx
from pathlib import Path
from typing import Dict, Optional, Tuple
from enhanced_equation_handler import EquationProcessor
from document_creator import DocumentCreatorWithAnchors

class IntegratedMammothConverter:
    """Enhanced mammoth converter with fixed equation handling and anchor support."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.equation_processor = EquationProcessor()
        self.style_map = """
        p[style-name='Heading 1'] => h1:fresh
        p[style-name='Heading 2'] => h2:fresh
        p[style-name='Heading 3'] => h3:fresh
        p[style-name='Title'] => h1.title:fresh
        p[style-name='Subtitle'] => h2.subtitle:fresh
        r[style-name='Strong'] => strong
        r[style-name='Emphasis'] => em
        """
        
    def convert_with_equation_fix(self, docx_path: Path, output_path: Path) -> Tuple[str, Dict]:
        """Convert document with proper equation handling and anchor generation."""
        
        self.logger.info(f"Converting {docx_path.name} with enhanced equation handling...")
        
        # Step 1: Extract all equations first
        equations = self.equation_processor.extract_all_equations(docx_path)
        total_equations = (len(equations['office_math']) + 
                          len(equations['latex']) + 
                          len(equations['images']))
        
        self.logger.info(f"Found {total_equations} equations total:")
        self.logger.info(f"  - Office Math: {len(equations['office_math'])}")
        self.logger.info(f"  - LaTeX: {len(equations['latex'])}")
        self.logger.info(f"  - Equation Images: {len(equations['images'])}")
        
        # Step 2: Pre-process document for equation markers
        processed_doc_path = self._preprocess_document(docx_path, equations)
        
        # Step 3: Convert with mammoth
        html_content = self._convert_with_mammoth(processed_doc_path)
        
        # Step 4: Post-process HTML to fix equations
        html_content = self._postprocess_equations(html_content, equations)
        
        # Step 5: Add anchors to equations and images
        html_content = self.equation_processor.create_html_with_anchors(html_content, equations)
        
        # Step 6: Build complete HTML document
        complete_html = self._build_enhanced_html(html_content, equations)
        
        # Step 7: Save HTML
        output_path.write_text(complete_html, encoding='utf-8')
        self.logger.info(f"Saved converted document to {output_path}")
        
        # Return HTML and equation data for further processing
        return complete_html, equations
    
    def _preprocess_document(self, docx_path: Path, equations: Dict) -> Path:
        """Pre-process document to add equation markers."""
        import tempfile
        import shutil
        
        # Create temporary copy
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        shutil.copy2(docx_path, temp_file.name)
        temp_path = Path(temp_file.name)
        
        try:
            doc = docx.Document(temp_path)
            
            # Process each paragraph
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                
                # Check for Office Math equations in this paragraph
                para_equations = [eq for eq in equations['office_math'] 
                                 if eq['paragraph'] == para_idx + 1]
                
                if para_equations:
                    # Clear paragraph and rebuild with markers
                    new_runs = []
                    
                    for eq in para_equations:
                        marker = f"[EQ_MARKER_{eq['id']}]"
                        new_runs.append(marker)
                    
                    # Clear existing runs
                    for run in paragraph.runs:
                        run.text = ""
                    
                    # Add text with markers
                    if para_text.strip():
                        paragraph.add_run(para_text + " ")
                    
                    for marker in new_runs:
                        paragraph.add_run(marker + " ")
            
            doc.save(temp_path)
            
        except Exception as e:
            self.logger.error(f"Error in preprocessing: {e}")
        
        return temp_path
    
    def _convert_with_mammoth(self, docx_path: Path) -> str:
        """Convert document with mammoth."""
        try:
            with open(docx_path, "rb") as docx_file:
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=self.style_map,
                    ignore_empty_paragraphs=False
                )
            
            # Log warnings but filter out equation-related ones
            if result.messages:
                for msg in result.messages:
                    msg_str = str(msg)
                    if not any(term in msg_str.lower() for term in ['omath', 'equation', 'formula']):
                        self.logger.warning(f"Mammoth warning: {msg}")
            
            return result.value
            
        except Exception as e:
            self.logger.error(f"Error in mammoth conversion: {e}")
            return ""
    
    def _postprocess_equations(self, html_content: str, equations: Dict) -> str:
        """Post-process HTML to properly format equations."""
        
        # Replace Office Math markers
        for eq in equations['office_math']:
            marker = f"[EQ_MARKER_{eq['id']}]"
            
            if eq['type'] == 'display':
                replacement = f'<div class="equation display-math" data-anchor="{eq["anchor"]}">$${eq["latex"]}$$</div>'
            else:
                replacement = f'<span class="equation inline-math" data-anchor="{eq["anchor"]}">${eq["latex"]}$</span>'
            
            html_content = html_content.replace(marker, replacement)
        
        # Fix LaTeX equations that might have been escaped
        html_content = self._fix_latex_escaping(html_content)
        
        # Ensure proper spacing around equations
        html_content = re.sub(r'([.!?])\s*(\$\$[^$]+\$\$)', r'\1</p><p>\2', html_content)
        html_content = re.sub(r'(\$\$[^$]+\$\$)\s*([A-Z])', r'\1</p><p>\2', html_content)
        
        return html_content
    
    def _fix_latex_escaping(self, html: str) -> str:
        """Fix common LaTeX escaping issues."""
        fixes = [
            # Fix escaped dollar signs
            (r'\\\$', '$'),
            # Fix double backslashes
            (r'\\\\([a-zA-Z])', r'\\\1'),
            # Fix HTML entities in equations
            (r'&lt;', '<'),
            (r'&gt;', '>'),
            (r'&amp;', '&'),
            # Fix spaces in equations
            (r'\$\s+([^$]+?)\s+\$', r'$\1$'),
            (r'\$\$\s+([^$]+?)\s+\$\$', r'$$\1$$'),
        ]
        
        for pattern, replacement in fixes:
            html = re.sub(pattern, replacement, html)
        
        # Special handling for equation blocks
        def fix_equation_block(match):
            eq_content = match.group(1)
            # Remove extra spaces and line breaks
            eq_content = ' '.join(eq_content.split())
            return f'$${eq_content}$$'
        
        html = re.sub(r'\$\$(.*?)\$\$', fix_equation_block, html, flags=re.DOTALL)
        
        return html
    
    def _build_enhanced_html(self, body_html: str, equations: Dict) -> str:
        """Build complete HTML with enhanced equation support."""
        
        # Determine if we need MathJax
        has_equations = any(len(equations[key]) > 0 for key in equations)
        
        # MathJax configuration
        mathjax_config = """
    <script>
        window.MathJax = {
            tex: {
                inlineMath: [['$', '$'], ['\\\\(', '\\\\)']],
                displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']],
                processEscapes: true,
                processEnvironments: true,
                processRefs: true
            },
            svg: {
                fontCache: 'global',
                mtextInheritFont: true,
                merrorInheritFont: true,
                mtextFont: '',
                merrorFont: 'serif',
                scale: 1.1
            },
            options: {
                renderActions: {
                    addMenu: [0, '', ''],
                    checkLoading: [0, '', '']
                }
            }
        };
    </script>
    <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
""" if has_equations else ""
        
        # Enhanced CSS for equations and anchors
        enhanced_css = """
        /* Equation styles */
        .equation {
            margin: 0.5em 0;
            position: relative;
        }
        
        .display-math {
            display: block;
            text-align: center;
            margin: 1em 0;
            padding: 0.5em;
            overflow-x: auto;
        }
        
        .inline-math {
            display: inline;
            padding: 0 0.2em;
        }
        
        /* Anchor styles */
        .equation-anchor {
            position: absolute;
            left: -30px;
            top: 50%;
            transform: translateY(-50%);
            width: 20px;
            height: 20px;
            opacity: 0;
        }
        
        .equation:hover .equation-anchor {
            opacity: 0.3;
        }
        
        /* Equation numbering */
        .equation-number {
            float: right;
            margin-right: 1em;
            color: #666;
        }
        
        /* Error handling for failed equations */
        .equation-error {
            color: red;
            border: 1px solid red;
            padding: 0.5em;
            background: #ffe6e6;
            font-family: monospace;
        }
        
        /* Responsive equations */
        @media screen and (max-width: 600px) {
            .display-math {
                font-size: 0.9em;
                padding: 0.3em;
            }
        }
        
        /* Print styles for equations */
        @media print {
            .equation-anchor {
                display: none !important;
            }
            
            .display-math {
                page-break-inside: avoid;
            }
        }
"""
        
        return f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document with Enhanced Equations</title>
    {mathjax_config}
    <style>
        body {{
            font-family: 'Amiri', 'Arial', sans-serif;
            line-height: 1.8;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            direction: rtl;
            text-align: right;
            color: #333;
        }}
        
        h1, h2, h3, h4 {{
            color: #1a1a1a;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        
        .title {{
            text-align: center;
            font-size: 2.5em;
            margin-bottom: 0.2em;
            color: #0066cc;
        }}
        
        {enhanced_css}
        
        /* Table styles */
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        
        td, th {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: right;
        }}
        
        th {{
            background-color: #f5f5f5;
            font-weight: bold;
        }}
        
        /* Image styles */
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1em auto;
        }}
    </style>
</head>
<body>
    <div class="content">
        {body_html}
    </div>
    
    <!-- Equation error handler -->
    <script>
        // Handle MathJax errors
        window.addEventListener('load', function() {{
            if (window.MathJax) {{
                MathJax.startup.document.addEventListener('math error', function(e) {{
                    console.error('MathJax error:', e);
                    const el = e.target;
                    el.classList.add('equation-error');
                    el.title = 'Error rendering equation: ' + e.message;
                }});
            }}
        }});
        
        // Add equation numbering
        document.addEventListener('DOMContentLoaded', function() {{
            const displayEquations = document.querySelectorAll('.display-math');
            displayEquations.forEach((eq, index) => {{
                const number = document.createElement('span');
                number.className = 'equation-number';
                number.textContent = `({{index + 1}})`;
                eq.appendChild(number);
            }});
        }});
    </script>
</body>
</html>"""
    
    def create_word_with_anchors(self, content_data: Dict, output_path: Path):
        """Create a Word document with anchors from converted content."""
        creator = DocumentCreatorWithAnchors()
        
        doc = creator.create_document(
            title=content_data.get('title', 'Converted Document'),
            author=content_data.get('author', 'Conversion System')
        )
        
        # Add content with anchors
        for item in content_data.get('content', []):
            if item['type'] == 'paragraph_with_equation':
                creator.add_paragraph_with_equation(
                    text=item['text'],
                    equation=item['equation'],
                    equation_type=item.get('equation_type', 'latex'),
                    position=item.get('position', 'inline')
                )
            elif item['type'] == 'image':
                creator.add_image_with_anchor(
                    image_path=item['path'],
                    width=item.get('width'),
                    caption=item.get('caption', ''),
                    alt_text=item.get('alt_text', '')
                )
            elif item['type'] == 'table':
                creator.add_table_with_equations(
                    data=item['data'],
                    has_header=item.get('has_header', True)
                )
        
        # Save document
        creator.save_document(output_path)
        
        return creator.generate_anchor_report()


# Example usage and testing
if __name__ == "__main__":
    import logging
    
    # Setup logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Initialize converter
    converter = IntegratedMammothConverter()
    
    # Convert a document
    input_doc = Path("input/sample_with_equations.docx")
    output_html = Path("output/converted_with_anchors.html")
    
    # Ensure output directory exists
    output_html.parent.mkdir(parents=True, exist_ok=True)
    
    # Convert with equation fixes
    if input_doc.exists():
        html_content, equations = converter.convert_with_equation_fix(input_doc, output_html)
        
        print(f"\nConversion complete!")
        print(f"Output saved to: {output_html}")
        print(f"\nEquation Summary:")
        print(f"  Office Math: {len(equations['office_math'])}")
        print(f"  LaTeX: {len(equations['latex'])}")
        print(f"  Images: {len(equations['images'])}")
        
        # Create Word document with anchors
        word_output = Path("output/document_with_anchors.docx")
        
        # Prepare content data
        content_data = {
            'title': 'Document with Equation Anchors',
            'author': 'Conversion System',
            'content': [
                {
                    'type': 'paragraph_with_equation',
                    'text': 'Example equation from conversion:',
                    'equation': 'E = mc^2',
                    'equation_type': 'latex',
                    'position': 'inline'
                }
            ]
        }
        
        report = converter.create_word_with_anchors(content_data, word_output)
        print(f"\nWord document created: {word_output}")
        print(report)
    else:
        print(f"Input file not found: {input_doc}")
