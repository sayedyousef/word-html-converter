# enhanced_equation_handler.py
"""Enhanced equation handling for Word to HTML conversion."""

import re
import logging
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, List, Tuple
from pathlib import Path

class EquationProcessor:
    """Advanced processor for handling equations in Word documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.equation_counter = 0
        self.equation_registry = {}
    
    def extract_all_equations(self, docx_path: Path) -> Dict[str, any]:
        """Extract all types of equations from Word document."""
        equations = {
            'office_math': [],
            'latex': [],
            'images': []
        }
        
        # Extract Office Math equations with positions
        equations['office_math'] = self._extract_office_math_with_positions(docx_path)
        
        # Extract LaTeX equations
        equations['latex'] = self._extract_latex_equations(docx_path)
        
        # Extract equation images
        equations['images'] = self._extract_equation_images(docx_path)
        
        return equations
    
    def _extract_office_math_with_positions(self, docx_path: Path) -> List[Dict]:
        """Extract Office Math equations with their document positions."""
        equations = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                if 'word/document.xml' not in zip_file.namelist():
                    return equations
                
                with zip_file.open('word/document.xml') as xml_file:
                    tree = ET.parse(xml_file)
                    root = tree.getroot()
                    
                    # Namespaces for parsing
                    ns = {
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                    }
                    
                    # Track paragraph index for position mapping
                    para_index = 0
                    
                    # Find all paragraphs
                    for para in root.findall('.//w:p', ns):
                        para_index += 1
                        
                        # Check for Office Math in paragraph
                        math_elements = para.findall('.//m:oMath', ns)
                        
                        for math_idx, math_elem in enumerate(math_elements):
                            eq_id = f"eq_{para_index}_{math_idx}"
                            
                            # Extract equation content
                            equation_data = self._parse_office_math_element(math_elem, ns)
                            
                            equations.append({
                                'id': eq_id,
                                'paragraph': para_index,
                                'position': math_idx,
                                'type': 'inline' if self._is_inline_math(math_elem, ns) else 'display',
                                'content': equation_data['text'],
                                'latex': equation_data['latex'],
                                'anchor': f"equation-{eq_id}"
                            })
                    
        except Exception as e:
            self.logger.error(f"Error extracting Office Math: {e}")
        
        return equations
    
    def _parse_office_math_element(self, math_elem, ns: Dict) -> Dict:
        """Parse Office Math element to extract content and convert to LaTeX."""
        text_parts = []
        latex_parts = []
        
        # Process different math structures
        for elem in math_elem.iter():
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            
            if tag == 't':  # Text element
                if elem.text:
                    text_parts.append(elem.text)
                    
            elif tag == 'frac':  # Fraction
                num = self._get_math_text(elem.find('.//m:num', ns), ns)
                den = self._get_math_text(elem.find('.//m:den', ns), ns)
                latex_parts.append(f"\\frac{{{num}}}{{{den}}}")
                
            elif tag == 'rad':  # Radical (square root)
                deg = elem.find('.//m:deg', ns)
                rad_content = self._get_math_text(elem.find('.//m:e', ns), ns)
                if deg is not None and self._get_math_text(deg, ns):
                    latex_parts.append(f"\\sqrt[{self._get_math_text(deg, ns)}]{{{rad_content}}}")
                else:
                    latex_parts.append(f"\\sqrt{{{rad_content}}}")
                    
            elif tag == 'sup':  # Superscript
                base = self._get_math_text(elem.find('.//m:e', ns), ns)
                sup = self._get_math_text(elem.find('.//m:sup', ns), ns)
                latex_parts.append(f"{base}^{{{sup}}}")
                
            elif tag == 'sub':  # Subscript
                base = self._get_math_text(elem.find('.//m:e', ns), ns)
                sub = self._get_math_text(elem.find('.//m:sub', ns), ns)
                latex_parts.append(f"{base}_{{{sub}}}")
        
        # Combine text and generate LaTeX
        text = ' '.join(text_parts)
        latex = ' '.join(latex_parts) if latex_parts else self._text_to_latex(text)
        
        return {
            'text': text,
            'latex': latex
        }
    
    def _get_math_text(self, elem, ns: Dict) -> str:
        """Get text content from math element."""
        if elem is None:
            return ""
        
        text_parts = []
        for t_elem in elem.findall('.//m:t', ns):
            if t_elem.text:
                text_parts.append(t_elem.text)
        
        return ''.join(text_parts)
    
    def _text_to_latex(self, text: str) -> str:
        """Convert plain text to LaTeX format with symbol replacements."""
        replacements = [
            ('÷', '\\div'), ('×', '\\times'), ('±', '\\pm'),
            ('≈', '\\approx'), ('≠', '\\neq'), ('≤', '\\leq'),
            ('≥', '\\geq'), ('∞', '\\infty'), ('∑', '\\sum'),
            ('∫', '\\int'), ('√', '\\sqrt'), ('∂', '\\partial'),
            ('∈', '\\in'), ('∉', '\\notin'), ('∅', '\\emptyset'),
            ('α', '\\alpha'), ('β', '\\beta'), ('γ', '\\gamma'),
            ('δ', '\\delta'), ('ε', '\\epsilon'), ('θ', '\\theta'),
            ('λ', '\\lambda'), ('μ', '\\mu'), ('π', '\\pi'),
            ('σ', '\\sigma'), ('τ', '\\tau'), ('φ', '\\phi'),
            ('ω', '\\omega'), ('Σ', '\\Sigma'), ('Δ', '\\Delta'),
            ('Ω', '\\Omega'), ('→', '\\rightarrow'), ('←', '\\leftarrow'),
            ('⇒', '\\Rightarrow'), ('⇔', '\\Leftrightarrow'),
        ]
        
        latex = text
        for old, new in replacements:
            latex = latex.replace(old, new)
        
        # Detect and convert fractions
        latex = re.sub(r'(\d+)\s*/\s*(\d+)', r'\\frac{\1}{\2}', latex)
        
        # Detect exponents
        latex = re.sub(r'(\w+)\^(\d+)', r'\1^{\2}', latex)
        latex = re.sub(r'(\w+)_(\d+)', r'\1_{\2}', latex)
        
        return latex
    
    def _is_inline_math(self, math_elem, ns: Dict) -> bool:
        """Check if math element is inline or display."""
        parent = math_elem.getparent()
        if parent is not None and parent.tag.endswith('oMathPara'):
            return False
        return True
    
    def _extract_latex_equations(self, docx_path: Path) -> List[Dict]:
        """Extract LaTeX equations with positions."""
        equations = []
        
        try:
            import mammoth
            with open(docx_path, "rb") as f:
                raw_result = mammoth.extract_raw_text(f)
                raw_text = raw_result.value
            
            # Find all LaTeX patterns with positions
            patterns = [
                (r'\$\$([^$]+)\$\$', 'display'),
                (r'\$([^$\n]+)\$', 'inline'),
                (r'\\\[([^\]]+)\\\]', 'display'),
                (r'\\\(([^\)]+)\\\)', 'inline')
            ]
            
            for pattern, eq_type in patterns:
                for match in re.finditer(pattern, raw_text):
                    equations.append({
                        'id': f"latex_{len(equations)}",
                        'type': eq_type,
                        'latex': match.group(0),
                        'content': match.group(1),
                        'position': match.start(),
                        'anchor': f"equation-latex-{len(equations)}"
                    })
            
        except Exception as e:
            self.logger.debug(f"Error extracting LaTeX: {e}")
        
        return equations
    
    def _extract_equation_images(self, docx_path: Path) -> List[Dict]:
        """Extract equation images (for equations inserted as images)."""
        equation_images = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # Check relationships for equation images
                if 'word/_rels/document.xml.rels' in zip_file.namelist():
                    with zip_file.open('word/_rels/document.xml.rels') as rels_file:
                        tree = ET.parse(rels_file)
                        root = tree.getroot()
                        
                        ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                        
                        for rel in root.findall('.//r:Relationship', ns):
                            target = rel.get('Target')
                            if target and 'media/' in target:
                                # Check if image might be an equation
                                if any(kw in target.lower() for kw in ['equation', 'eq', 'formula']):
                                    equation_images.append({
                                        'id': f"img_eq_{len(equation_images)}",
                                        'path': target,
                                        'anchor': f"equation-img-{len(equation_images)}"
                                    })
        
        except Exception as e:
            self.logger.debug(f"Error extracting equation images: {e}")
        
        return equation_images
    
    def create_html_with_anchors(self, html_content: str, equations: Dict) -> str:
        """Add equation anchors to HTML content."""
        modified_html = html_content
        
        # Add anchors for Office Math equations
        for eq in equations.get('office_math', []):
            anchor = f'<a id="{eq["anchor"]}" class="equation-anchor"></a>'
            
            # Try to insert anchor near equation content
            if eq['type'] == 'display':
                # For display equations, wrap in div with anchor
                eq_html = f'{anchor}<div class="equation display-equation">$${eq["latex"]}$$</div>'
            else:
                # For inline equations
                eq_html = f'{anchor}<span class="equation inline-equation">${eq["latex"]}$</span>'
            
            # Replace equation markers if they exist
            marker = f"[EQUATION_{eq['id']}]"
            if marker in modified_html:
                modified_html = modified_html.replace(marker, eq_html)
        
        # Add anchors for LaTeX equations
        for eq in equations.get('latex', []):
            anchor = f'<a id="{eq["anchor"]}" class="equation-anchor"></a>'
            
            # Find and replace the LaTeX equation with anchor
            original = eq['latex']
            replacement = f'{anchor}{original}'
            modified_html = modified_html.replace(original, replacement, 1)
        
        return modified_html


# document_creator.py
"""Create Word documents with special anchors for equations and images."""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Optional
import re

class DocumentCreatorWithAnchors:
    """Create Word documents with special anchors for equations and images."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.doc = None
        self.anchor_registry = {}
        self.equation_counter = 0
        self.image_counter = 0
    
    def create_document(self, title: str = "Document", author: str = "Unknown") -> Document:
        """Create a new Word document with metadata."""
        self.doc = Document()
        
        # Set document properties
        self.doc.core_properties.title = title
        self.doc.core_properties.author = author
        
        # Add title to document
        title_para = self.doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add author
        author_para = self.doc.add_paragraph(f"Author: {author}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Empty line
        
        return self.doc
    
    def add_paragraph_with_equation(self, text: str, equation: str, 
                                   equation_type: str = 'latex',
                                   position: str = 'inline') -> str:
        """Add paragraph with equation and return anchor ID."""
        para = self.doc.add_paragraph()
        
        # Generate unique anchor ID
        self.equation_counter += 1
        anchor_id = f"eq-anchor-{self.equation_counter}"
        
        # Add text before equation
        if text:
            para.add_run(text + " ")
        
        # Add anchor marker (as bookmark)
        self._add_bookmark(para, anchor_id)
        
        # Add equation based on type
        if equation_type == 'latex':
            # Add LaTeX equation as text (will be processed later)
            if position == 'display':
                eq_text = f"$${equation}$$"
                eq_para = self.doc.add_paragraph(eq_text)
                eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                para.add_run(f"${equation}$")
        else:
            # For Office Math, add as placeholder
            para.add_run(f"[EQUATION: {equation}]")
        
        # Register anchor
        self.anchor_registry[anchor_id] = {
            'type': 'equation',
            'content': equation,
            'equation_type': equation_type,
            'position': position
        }
        
        return anchor_id
    
    def add_image_with_anchor(self, image_path: str, width: Optional[float] = None,
                             caption: str = "", alt_text: str = "") -> str:
        """Add image with anchor and return anchor ID."""
        # Generate unique anchor ID
        self.image_counter += 1
        anchor_id = f"img-anchor-{self.image_counter}"
        
        # Add anchor paragraph
        anchor_para = self.doc.add_paragraph()
        self._add_bookmark(anchor_para, anchor_id)
        
        # Add image
        try:
            if width:
                self.doc.add_picture(image_path, width=Inches(width))
            else:
                self.doc.add_picture(image_path)
            
            # Add caption if provided
            if caption:
                caption_para = self.doc.add_paragraph(caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.style = 'Caption'
        
        except Exception as e:
            self.logger.error(f"Error adding image: {e}")
            self.doc.add_paragraph(f"[IMAGE PLACEHOLDER: {image_path}]")
        
        # Register anchor
        self.anchor_registry[anchor_id] = {
            'type': 'image',
            'path': image_path,
            'caption': caption,
            'alt_text': alt_text or caption
        }
        
        return anchor_id
    
    def _add_bookmark(self, paragraph, bookmark_name: str):
        """Add bookmark (anchor) to paragraph."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Create bookmark start element
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(len(self.anchor_registry)))
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        # Create bookmark end element
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(len(self.anchor_registry)))
        
        # Add to paragraph
        paragraph._p.append(bookmark_start)
        paragraph._p.append(bookmark_end)
    
    def add_table_with_equations(self, data: List[List[str]], 
                                has_header: bool = True) -> List[str]:
        """Add table with potential equations in cells."""
        anchor_ids = []
        
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                
                # Check if cell contains equation
                if self._contains_equation(cell_text):
                    # Extract equation
                    equation = self._extract_equation(cell_text)
                    clean_text = cell_text.replace(equation, "")
                    
                    # Add with anchor
                    para = cell.paragraphs[0]
                    para.text = clean_text
                    
                    # Add anchor
                    self.equation_counter += 1
                    anchor_id = f"table-eq-{row_idx}-{col_idx}-{self.equation_counter}"
                    self._add_bookmark(para, anchor_id)
                    
                    # Add equation
                    para.add_run(equation)
                    
                    anchor_ids.append(anchor_id)
                    
                    # Register
                    self.anchor_registry[anchor_id] = {
                        'type': 'table_equation',
                        'content': equation,
                        'position': f"row_{row_idx}_col_{col_idx}"
                    }
                else:
                    cell.text = cell_text
                
                # Make header row bold
                if has_header and row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        
        return anchor_ids
    
    def _contains_equation(self, text: str) -> bool:
        """Check if text contains equation patterns."""
        patterns = [r'\$[^$]+\$', r'\$\$[^$]+\$\$', r'\\\[', r'\\\(']
        return any(re.search(pattern, text) for pattern in patterns)
    
    def _extract_equation(self, text: str) -> str:
        """Extract equation from text."""
        patterns = [
            (r'(\$\$[^$]+\$\$)', 1),
            (r'(\$[^$]+\$)', 1),
            (r'(\\\[[^\]]+\\\])', 1),
            (r'(\\\([^\)]+\\\))', 1)
        ]
        
        for pattern, group in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(group)
        
        return ""
    
    def add_references_with_links(self, references: List[Dict[str, str]]):
        """Add references section with links to anchors."""
        self.doc.add_heading('References to Equations and Images', 1)
        
        for anchor_id, info in self.anchor_registry.items():
            para = self.doc.add_paragraph()
            
            if info['type'] == 'equation':
                text = f"Equation {anchor_id}: {info['content'][:50]}..."
            elif info['type'] == 'image':
                text = f"Image {anchor_id}: {info['caption'] or info['path']}"
            else:
                text = f"Element {anchor_id}"
            
            # Add as hyperlink to bookmark
            self._add_internal_hyperlink(para, anchor_id, text)
    
    def _add_internal_hyperlink(self, paragraph, bookmark_name: str, text: str):
        """Add internal hyperlink to bookmark."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        # Create run with text
        run = OxmlElement('w:r')
        run_properties = OxmlElement('w:rPr')
        
        # Add blue color and underline
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        run_properties.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        run_properties.append(underline)
        
        run.append(run_properties)
        
        # Add text
        text_element = OxmlElement('w:t')
        text_element.text = text
        run.append(text_element)
        
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
    
    def save_document(self, output_path: Path):
        """Save document with anchor registry."""
        if self.doc:
            # Save document
            self.doc.save(output_path)
            
            # Save anchor registry as companion file
            import json
            registry_path = output_path.with_suffix('.anchors.json')
            with open(registry_path, 'w', encoding='utf-8') as f:
                json.dump(self.anchor_registry, f, indent=2, ensure_ascii=False)
            
            self.logger.info(f"Document saved to {output_path}")
            self.logger.info(f"Anchor registry saved to {registry_path}")
            self.logger.info(f"Total anchors created: {len(self.anchor_registry)}")
        else:
            self.logger.error("No document to save")
    
    def generate_anchor_report(self) -> str:
        """Generate a report of all anchors in the document."""
        report = []
        report.append("=" * 60)
        report.append("ANCHOR REGISTRY REPORT")
        report.append("=" * 60)
        
        # Group by type
        by_type = {}
        for anchor_id, info in self.anchor_registry.items():
            anchor_type = info['type']
            if anchor_type not in by_type:
                by_type[anchor_type] = []
            by_type[anchor_type].append((anchor_id, info))
        
        # Report by type
        for anchor_type, items in by_type.items():
            report.append(f"\n{anchor_type.upper()} ({len(items)} items):")
            report.append("-" * 40)
            
            for anchor_id, info in items:
                if anchor_type == 'equation':
                    report.append(f"  {anchor_id}: {info['equation_type']} - {info['position']}")
                    report.append(f"    Content: {info['content'][:50]}...")
                elif anchor_type == 'image':
                    report.append(f"  {anchor_id}: {info['path']}")
                    if info['caption']:
                        report.append(f"    Caption: {info['caption']}")
                else:
                    report.append(f"  {anchor_id}: {str(info)[:100]}...")
        
        report.append("\n" + "=" * 60)
        return "\n".join(report)


# Example usage
if __name__ == "__main__":
    # Setup logging
    logging.basicConfig(level=logging.INFO)
    
    # Create document with anchors
    creator = DocumentCreatorWithAnchors()
    doc = creator.create_document(
        title="Mathematical Document with Anchors",
        author="Document Processing System"
    )
    
    # Add content with equations
    anchor1 = creator.add_paragraph_with_equation(
        "The quadratic formula is",
        "x = \\frac{-b \\pm \\sqrt{b^2-4ac}}{2a}",
        equation_type='latex',
        position='display'
    )
    
    anchor2 = creator.add_paragraph_with_equation(
        "Einstein's famous equation",
        "E = mc^2",
        equation_type='latex',
        position='inline'
    )
    
    # Add image with anchor
    # anchor3 = creator.add_image_with_anchor(
    #     "path/to/image.png",
    #     width=4.0,
    #     caption="Sample diagram",
    #     alt_text="A sample diagram showing the process"
    # )
    
    # Add table with equations
    table_data = [
        ["Variable", "Formula", "Value"],
        ["Area", "$A = \\pi r^2$", "28.27"],
        ["Volume", "$V = \\frac{4}{3}\\pi r^3$", "113.1"],
    ]
    table_anchors = creator.add_table_with_equations(table_data)
    
    # Add references section
    creator.add_references_with_links([])
    
    # Save document
    output_path = Path("output/document_with_anchors.docx")
    output_path.parent.mkdir(exist_ok=True)
    creator.save_document(output_path)
    
    # Print report
    print(creator.generate_anchor_report())
