# word_anchor_adder.py - COMPLETE VERSION FOR BOTH CASES
"""Add anchors to Word documents - handles BOTH Office Math and LaTeX."""

import logging
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import zipfile
import xml.etree.ElementTree as ET
import re
from docx.shared import RGBColor


class WordAnchorAdder:
    """Add anchors to Word documents for both Office Math and LaTeX equations."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.anchor_counter = 0
        self.anchor_registry = {}
    
    def add_anchors_to_document(self, input_path: Path, output_path: Path) -> dict:
        """Add anchors to equations (both types) and images."""
        
        self.logger.info(f"Adding anchors to: {input_path.name}")
        self.anchor_counter = 0
        self.anchor_registry = {}
        
        # Open document
        doc = Document(input_path)
        
        # Method 1: Find Office Math equations
        office_math_paragraphs = self._find_office_math_paragraphs(input_path)
        
        # Process all paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            
            # Check if this paragraph has Office Math
            if para_idx in office_math_paragraphs:
                self.anchor_counter += 1
                anchor_id = f"eq-office-{self.anchor_counter}"
                self._add_bookmark_to_paragraph(paragraph, anchor_id)
                
                self.anchor_registry[anchor_id] = {
                    'type': 'office_math_equation',
                    'paragraph': para_idx
                }
                self.logger.debug(f"Added Office Math anchor {anchor_id}")
            
            # ALSO check for LaTeX patterns in text
            elif paragraph.text:
                text = paragraph.text
                # Check for LaTeX patterns
                latex_patterns = [
                    r'\$\$[^$]+\$\$',  # Display math
                    r'\$[^$\n]+\$'     # Inline math
                ]
                
                has_latex = False
                for pattern in latex_patterns:
                    if re.search(pattern, text):
                        has_latex = True
                        break
                
                if has_latex:
                    self.anchor_counter += 1
                    anchor_id = f"eq-latex-{self.anchor_counter}"
                    self._add_bookmark_to_paragraph(paragraph, anchor_id)
                    
                    self.anchor_registry[anchor_id] = {
                        'type': 'latex_equation',
                        'paragraph': para_idx,
                        'text_preview': text[:50]
                    }
                    self.logger.debug(f"Added LaTeX anchor {anchor_id}")
        
        # Process images
        self._add_image_anchors(doc)
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        # Save anchor registry
        import json
        registry_path = output_path.with_suffix('.anchors.json')
        with open(registry_path, 'w', encoding='utf-8') as f:
            json.dump(self.anchor_registry, f, indent=2, ensure_ascii=False)
        
        self.logger.info(f"Saved: {output_path}")
        self.logger.info(f"Total anchors: {len(self.anchor_registry)} (Office Math + LaTeX + Images)")
        
        return self.anchor_registry
    
    def _find_office_math_paragraphs(self, docx_path):
        """Find paragraphs containing Office Math equations."""
        equation_paragraphs = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                if 'word/document.xml' in zip_file.namelist():
                    with zip_file.open('word/document.xml') as xml_file:
                        tree = ET.parse(xml_file)
                        root = tree.getroot()
                        
                        ns = {
                            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                            'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                        }
                        
                        para_index = 0
                        for para in root.findall('.//w:p', ns):
                            # Check for Office Math
                            math_elements = para.findall('.//m:oMath', ns)
                            math_para_elements = para.findall('.//m:oMathPara', ns)
                            
                            if math_elements or math_para_elements:
                                equation_paragraphs.append(para_index)
                            
                            para_index += 1
                            
        except Exception as e:
            self.logger.warning(f"Could not parse Office Math: {e}")
        
        if equation_paragraphs:
            self.logger.info(f"Found {len(equation_paragraphs)} Office Math paragraphs")
        
        return equation_paragraphs
    
    def _add_image_anchors(self, doc):
        """Add anchors for images."""
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Check for images in paragraph
            if 'graphic' in paragraph._element.xml or 'picture' in paragraph._element.xml:
                self.anchor_counter += 1
                anchor_id = f"img-{self.anchor_counter}"
                
                self._add_bookmark_to_paragraph(paragraph, anchor_id)
                
                self.anchor_registry[anchor_id] = {
                    'type': 'image',
                    'paragraph': para_idx
                }
                self.logger.debug(f"Added image anchor {anchor_id}")
    
    def _add_bookmark_to_paragraph(self, paragraph, bookmark_name):
        """Add bookmark to paragraph."""
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(self.anchor_counter))
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(self.anchor_counter))
        
        paragraph._element.insert(0, bookmark_start)
        paragraph._element.append(bookmark_end)
    
def process_folder(self, input_folder: Path, output_folder: Path):
    """Process all documents in folder."""
    docx_files = list(input_folder.rglob("*.docx"))
    docx_files = [f for f in docx_files if not f.name.startswith("~")]
    
    self.logger.info(f"Processing {len(docx_files)} documents")
    
    for idx, docx_file in enumerate(docx_files, 1):
        # Use the SAME folder structure as HTML output
        safe_name = docx_file.stem.replace(' ', '_')
        article_folder = output_folder / f"article_{idx:03d}_{safe_name}"
        
        # Save anchored Word doc in the SAME folder as HTML
        output_path = article_folder / f"{docx_file.stem}_anchored.docx"
        
        self.add_anchors_to_document(docx_file, output_path)
        self.logger.info(f"[{idx}/{len(docx_files)}] Saved to {output_path}")
    
    return True