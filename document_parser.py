# document_parser.py
"""Parse Word documents and extract content."""

import logging
from pathlib import Path
from typing import Optional, List
import docx
from docx2python import docx2python
from models import DocumentContent, ImageInfo, FootnoteInfo
from utils import extract_text_safely, detect_latex_equations

class DocumentParser:
    """Handles parsing of Word documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse_document(self, file_path: Path) -> Optional[DocumentContent]:
        """Parse a Word document and extract all content."""
        try:
            self.logger.info(f"Parsing document: {file_path.name}")
            
            # Use python-docx for structure
            doc = docx.Document(file_path)
            
            # Use docx2python for footnotes
            with docx2python(str(file_path)) as docx_content:
                footnotes_data = docx_content.footnotes
            
            # Extract content
            content = DocumentContent(
                title=self._extract_title(doc),
                author=self._extract_author(doc),
                body_html="",
                footnotes=self._extract_footnotes(footnotes_data),
                images=[]
            )
            
            # Parse body content
            body_parts = []
            image_counter = 0
            
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # Handle paragraph
                    para = docx.text.paragraph.Paragraph(element, doc)
                    html = self._paragraph_to_html(para)
                    if html:
                        body_parts.append(html)
                        
                elif element.tag.endswith('tbl'):
                    # Handle table
                    table = docx.table.Table(element, doc)
                    html = self._table_to_html(table)
                    if html:
                        body_parts.append(html)
                
                # Check for images in runs
                for para in doc.paragraphs:
                    for run in para.runs:
                        if 'graphic' in run._element.xml:
                            image_counter += 1
                            img_info = ImageInfo(number=image_counter)
                            content.images.append(img_info)
            
            content.body_html = "\n".join(body_parts)
            
            # Check for equations
            has_eq, _ = detect_latex_equations(content.body_html)
            content.has_equations = has_eq
            
            return content
            
        except Exception as e:
            self.logger.error(f"Error parsing {file_path}: {e}")
            return None
    
    def _extract_title(self, doc) -> str:
        """Extract document title."""
        # Try core properties first
        if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
            return doc.core_properties.title
        
        # Fall back to first heading or paragraph
        for para in doc.paragraphs[:5]:  # Check first 5 paragraphs
            if para.style.name.startswith('Heading') or para.style.name == 'Title':
                return para.text.strip()
        
        # Use first non-empty paragraph
        for para in doc.paragraphs[:3]:
            if para.text.strip():
                return para.text.strip()[:100]
        
        return "Untitled"
    
    def _extract_author(self, doc) -> str:
        """Extract document author."""
        if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
            return doc.core_properties.author
        return "Unknown"
    

    def _extract_footnotes(self, footnotes_data) -> List[FootnoteInfo]:
        """Extract footnotes from docx2python data."""
        footnotes = []
        
        try:
            if footnotes_data:
                # docx2python returns footnotes as a flat list
                # Each footnote is a list of text runs
                for idx, footnote_content in enumerate(footnotes_data):
                    if footnote_content and isinstance(footnote_content, list):
                        # Join all text runs in the footnote
                        text_parts = []
                        for part in footnote_content:
                            if isinstance(part, list):
                                text_parts.extend(part)
                            else:
                                text_parts.append(str(part))
                        
                        text = " ".join(str(p) for p in text_parts if p)
                        
                        if text.strip():
                            has_latex, _ = detect_latex_equations(text)
                            
                            footnotes.append(FootnoteInfo(
                                id=str(idx + 1),
                                text=text,
                                contains_latex=has_latex
                            ))
                            
            self.logger.info(f"Extracted {len(footnotes)} footnotes")
            
        except Exception as e:
            self.logger.warning(f"Error extracting footnotes: {e}")
            self.logger.debug(f"Footnotes data structure: {type(footnotes_data)}")
        
        return footnotes

    def _paragraph_to_html(self, paragraph) -> str:
        """Convert paragraph to HTML."""
        text = extract_text_safely(paragraph)
        if not text:
            return ""
        
        # Determine HTML tag based on style
        style_name = paragraph.style.name
        
        if style_name.startswith('Heading 1'):
            return f"<h1>{text}</h1>"
        elif style_name.startswith('Heading 2'):
            return f"<h2>{text}</h2>"
        elif style_name.startswith('Heading 3'):
            return f"<h3>{text}</h3>"
        elif style_name.startswith('Heading 4'):
            return f"<h4>{text}</h4>"
        elif style_name == 'Title':
            return f"<h1 class='title'>{text}</h1>"
        else:
            # Handle inline formatting
            html_text = text
            
            # Check for bold, italic, etc.
            for run in paragraph.runs:
                run_text = run.text
                if run.bold:
                    html_text = html_text.replace(run_text, f"<strong>{run_text}</strong>")
                if run.italic:
                    html_text = html_text.replace(run_text, f"<em>{run_text}</em>")
            
            return f"<p>{html_text}</p>"
    
    def _table_to_html(self, table) -> str:
        """Convert table to HTML."""
        html_parts = ["<table class='document-table'>"]
        
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                cell_text = extract_text_safely(cell)
                html_parts.append(f"<td>{cell_text}</td>")
            html_parts.append("</tr>")
        
        html_parts.append("</table>")
        return "\n".join(html_parts)
