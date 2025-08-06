# anchor_generator.py
"""Generate Word documents with anchors from converted HTML."""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
from typing import Dict, List
import json

class AnchorGenerator:
    """Generate Word documents with anchors from converted HTML."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.anchor_registry = {}
    
    def create_from_html_data(self, html_path: Path, anchors_json_path: Path) -> Path:
        """Create Word document from HTML conversion data."""
        
        # Load anchor registry if exists
        if anchors_json_path.exists():
            with open(anchors_json_path, 'r', encoding='utf-8') as f:
                self.anchor_registry = json.load(f)
            self.logger.info(f"Loaded {len(self.anchor_registry)} anchors from {anchors_json_path.name}")
        else:
            self.logger.warning(f"No anchor registry found at {anchors_json_path}")
            return None
        
        # Create new document
        doc = Document()
        
        # Extract title from HTML filename
        title = html_path.stem.replace('_', ' ').title()
        
        # Add title
        doc.add_heading(f'{title} - With Anchors', 0)
        
        # Add content with anchors
        self._add_anchored_content(doc)
        
        # Save document
        output_path = html_path.parent / f"{html_path.stem}_anchored.docx"
        doc.save(output_path)
        
        self.logger.info(f"Created anchored document: {output_path}")
        return output_path
    
    def _add_anchored_content(self, doc):
        """Add content with anchors to document."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Add summary section
        doc.add_heading('Document Anchors Summary', 1)
        
        # Count anchors by type
        anchor_types = {}
        for anchor_id, info in self.anchor_registry.items():
            anchor_type = info.get('type', 'unknown')
            if anchor_type not in anchor_types:
                anchor_types[anchor_type] = 0
            anchor_types[anchor_type] += 1
        
        # Add summary paragraph
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total anchors: {len(self.anchor_registry)}\n")
        for anchor_type, count in anchor_types.items():
            summary_para.add_run(f"  • {anchor_type}: {count}\n")
        
        # Add detailed anchor index
        doc.add_heading('Anchor Index', 1)
        
        # Group anchors by type
        by_type = {}
        for anchor_id, info in self.anchor_registry.items():
            anchor_type = info.get('type', 'unknown')
            if anchor_type not in by_type:
                by_type[anchor_type] = []
            by_type[anchor_type].append((anchor_id, info))
        
        # Add sections for each type
        for anchor_type, anchors in by_type.items():
            doc.add_heading(f'{anchor_type.replace("_", " ").title()}s', 2)
            
            for anchor_id, info in anchors:
                para = doc.add_paragraph()
                
                # Add bookmark for this anchor
                bookmark_start = OxmlElement('w:bookmarkStart')
                bookmark_start.set(qn('w:id'), str(hash(anchor_id) % 100000))
                bookmark_start.set(qn('w:name'), anchor_id)
                
                bookmark_end = OxmlElement('w:bookmarkEnd')
                bookmark_end.set(qn('w:id'), str(hash(anchor_id) % 100000))
                
                para._p.append(bookmark_start)
                
                # Add content based on type
                if anchor_type == 'equation' or 'equation' in anchor_type:
                    content = info.get('content', info.get('latex', ''))
                    if len(content) > 50:
                        content = content[:50] + '...'
                    para.add_run(f"📐 {anchor_id}: {content}")
                elif anchor_type == 'image':
                    filename = info.get('filename', info.get('path', 'unknown'))
                    para.add_run(f"🖼️ {anchor_id}: {filename}")
                elif anchor_type == 'footnote':
                    number = info.get('number', '?')
                    para.add_run(f"📝 Footnote {number} (ID: {anchor_id})")
                else:
                    para.add_run(f"📌 {anchor_id}: {anchor_type}")
                
                para._p.append(bookmark_end)
        
        # Add navigation instructions
        doc.add_page_break()
        doc.add_heading('How to Use Anchors', 1)
        instructions = doc.add_paragraph()
        instructions.add_run(
            "These anchors can be used to create cross-references within the document:\n\n"
            "1. In Word, go to Insert > Link > Bookmark\n"
            "2. Select the anchor name from the list\n"
            "3. Create a cross-reference to link to it\n\n"
            "The anchors preserve the structure from the HTML conversion, "
            "allowing you to maintain references between equations, images, and footnotes."
        )

