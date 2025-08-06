# office_math_to_latex_converter.py
"""Convert Office Math equations to LaTeX in Word documents."""

import logging
from pathlib import Path
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

class OfficeMathToLatexConverter:
    """Convert Office Math to LaTeX text in Word documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.equation_count = 0
    
    def convert_document(self, input_path: Path, output_path: Path):
        """Convert Office Math to LaTeX in document."""
        
        self.logger.info(f"Converting Office Math to LaTeX in: {input_path.name}")
        
        # Find Office Math equations and their content
        equations = self._extract_office_math_as_text(input_path)
        
        if not equations:
            self.logger.info("No Office Math equations found")
            # Just copy the document
            import shutil
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return output_path
        
        # Open document
        doc = Document(input_path)
        
        # Replace equations with LaTeX
        para_idx = 0
        for paragraph in doc.paragraphs:
            if para_idx in equations:
                # Add LaTeX version
                eq_data = equations[para_idx]
                latex_text = f"$${eq_data['latex']}$$"  # Display equation
                
                # Clear paragraph and add text with LaTeX
                paragraph.clear()
                paragraph.add_run(eq_data.get('before_text', ''))
                paragraph.add_run(latex_text)
                paragraph.add_run(eq_data.get('after_text', ''))
                
                self.logger.debug(f"Replaced equation in paragraph {para_idx}")
            
            para_idx += 1
        
        # Save modified document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        self.logger.info(f"Saved LaTeX version to: {output_path}")
        self.logger.info(f"Converted {len(equations)} Office Math equations to LaTeX")
        
        return output_path
    
    def _extract_office_math_as_text(self, docx_path):
        """Extract Office Math and convert to simple LaTeX."""
        equations = {}
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                if 'word/document.xml' in zip_file.namelist():
                    with zip_file.open('word/document.xml') as xml_file:
                        tree = ET.parse(xml_file)
                        root = tree.getroot()
                        
                        ns = {
                            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                            'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                        }
                        
                        para_idx = 0
                        for para in root.findall('.//w:p', ns):
                            math_elements = para.findall('.//m:oMath', ns)
                            
                            if math_elements:
                                # Extract text from Office Math
                                for math in math_elements:
                                    latex = self._office_math_to_latex(math, ns)
                                    
                                    equations[para_idx] = {
                                        'latex': latex,
                                        'before_text': '',  # Text before equation
                                        'after_text': ''     # Text after equation
                                    }
                            
                            para_idx += 1
        
        except Exception as e:
            self.logger.error(f"Error extracting Office Math: {e}")
        
        return equations
    
    def _office_math_to_latex(self, math_elem, ns):
        """Convert Office Math XML to LaTeX (simplified)."""
        
        # Extract all text from the equation
        text_parts = []
        for t_elem in math_elem.findall('.//m:t', ns):
            if t_elem.text:
                text_parts.append(t_elem.text)
        
        equation_text = ''.join(text_parts)
        
        # Try to identify common patterns and convert
        latex = equation_text
        
        # Check for fractions
        frac_elems = math_elem.findall('.//m:frac', ns)
        if frac_elems:
            for frac in frac_elems:
                num = ''.join([t.text for t in frac.findall('.//m:num//m:t', ns) if t.text])
                den = ''.join([t.text for t in frac.findall('.//m:den//m:t', ns) if t.text])
                if num and den:
                    latex = f"\\frac{{{num}}}{{{den}}}"
        
        # Check for square roots
        rad_elems = math_elem.findall('.//m:rad', ns)
        if rad_elems:
            for rad in rad_elems:
                content = ''.join([t.text for t in rad.findall('.//m:e//m:t', ns) if t.text])
                if content:
                    latex = f"\\sqrt{{{content}}}"
        
        # Basic symbol replacements
        replacements = [
            ('±', '\\pm'), ('×', '\\times'), ('÷', '\\div'),
            ('≤', '\\leq'), ('≥', '\\geq'), ('≠', '\\neq'),
            ('∑', '\\sum'), ('∫', '\\int'), ('√', '\\sqrt'),
            ('α', '\\alpha'), ('β', '\\beta'), ('π', '\\pi')
        ]
        
        for old, new in replacements:
            latex = latex.replace(old, new)
        
        return latex


# Updated workflow in main3.py:
def convert_with_office_math_support():
    """Complete workflow with Office Math support."""
    
    # Step 1: Convert Office Math to LaTeX
    print("Step 1: Converting Office Math to LaTeX...")
    from office_math_to_latex_converter import OfficeMathToLatexConverter
    
    math_converter = OfficeMathToLatexConverter()
    temp_folder = Path("temp_latex_docs")
    temp_folder.mkdir(exist_ok=True)
    
    for docx_file in Config.INPUT_FOLDER.glob("*.docx"):
        if not docx_file.name.startswith("~"):
            latex_version = temp_folder / docx_file.name
            math_converter.convert_document(docx_file, latex_version)
    
    # Step 2: Add anchors to LaTeX version
    print("Step 2: Adding anchors...")
    from word_anchor_adder import WordAnchorAdder
    anchor_adder = WordAnchorAdder()
    anchor_adder.process_folder(temp_folder, Config.OUTPUT_FOLDER)
    
    # Step 3: Convert LaTeX version to HTML
    print("Step 3: Converting to HTML...")
    converter = MammothConverter()
    converter.convert_folder(temp_folder, Config.OUTPUT_FOLDER)
    
    print("Complete! Office Math → LaTeX → HTML with anchors")